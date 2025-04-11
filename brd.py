import streamlit as st
import pandas as pd
import os
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import groq
 
# === Initialize Groq client ===
client = groq.Client(api_key="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")
 
# === System Prompts per BRD Section ===
brd_sections = {
    "Executive Summary": "Act as a strategic business consultant...",
    "Business Objectives": "Act as a business user...",
    "Scope": "Act as a project manager of an implementation partner...",
    "Technical Data Flow": "Act as a software architect...",
    "Data Sources": "Act as a data engineer...",
    "Business Rules": "Act as a domain expert...",
}
 
# === Utility Functions ===
def extract_text_from_pdf(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    return "\n".join([page.extract_text() for page in pdf_reader.pages])
 
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    return "\n".join([para.text for para in doc.paragraphs])
 
def generate_section_questions(section, system_prompt, business_requirement):
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"""Generate 30–40 unique, non-repetitive, and well-structured questions for the '{section}' section of a Business Requirements Document.
Format the output as a numbered list. The questions should be tailored to extract clear and complete stakeholder input for BRD creation,
based on the following requirements:
 
{business_requirement}"""}
    ]
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=messages,
        temperature=0.2,
        max_tokens=4096
    )
    return response.choices[0].message.content.split("\n")
 
def save_questionnaire_to_excel(questions_by_section):
    rows = []
    for section, questions in questions_by_section.items():
        for q in questions:
            if q.strip():
                rows.append({"Section": section, "Questions": q.strip(), "Answers": ""})
    df = pd.DataFrame(rows)
    output = BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)
    return output
 
def generate_brd_section(section, qna):
    messages = [
        {"role": "system", "content": f"""You are a professional Business Analyst writing content for a Business Requirements Document (BRD).
Use a formal and concise tone. Ensure clarity, no repetition, and structure the output with bullet points."""},
        {"role": "user", "content": f"""Using the following Q&A, write a well-organized and professional '{section}' section for the BRD.
Summarize insights, combine related points, and avoid repeating similar answers:
 
{qna}"""}
    ]
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=messages,
        temperature=0.2,
        max_tokens=4096
    )
    return response.choices[0].message.content
 
def save_brd_to_word(brd_sections_content):
    doc = Document()
    doc.add_heading("Business Requirement Document (BRD)", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\n")
    for section, content in brd_sections_content.items():
        doc.add_heading(section, level=2)
        for line in content.strip().split("\n"):
            if line.strip():
                doc.add_paragraph(f"• {line.strip()}")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
