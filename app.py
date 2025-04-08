import streamlit as st
import pandas as pd
import os
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import groq
 
# === Initialize Groq client ===
client = groq.Client(api_key="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")
 
# === System Prompts per BRD Section ===
brd_sections = {
    "Executive Summary": "Act as a strategic business consultant...",
    "Business Objectives": "Act as a business user...",
    "Scope": "Act as a project manager of an implementation partner...",
    "Technical Data Flow": "Act as a software architect...",
    "Data Sources": "Act as a data engineer...",
    "Business Rules": "Act as a domain expert...",
}
 
# === Utility Functions ===
def extract_text_from_pdf(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    return "\n".join([page.extract_text() for page in pdf_reader.pages])
 
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    return "\n".join([para.text for para in doc.paragraphs])
 
def generate_section_questions(section, system_prompt, business_requirement):
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"""Generate 30–40 unique, non-repetitive, and well-structured questions for the '{section}' section of a Business Requirements Document.
Format the output as a numbered list. The questions should be tailored to extract clear and complete stakeholder input for BRD creation,
based on the following requirements:
 
{business_requirement}"""}
    ]
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=messages,
        temperature=0.2,
        max_tokens=4096
    )
    return response.choices[0].message.content.split("\n")
 
def save_questionnaire_to_excel(questions_by_section):
    rows = []
    for section, questions in questions_by_section.items():
        for q in questions:
            if q.strip():
                rows.append({"Section": section, "Questions": q.strip(), "Answers": ""})
    df = pd.DataFrame(rows)
    output = BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)
    return output
 
def generate_brd_section(section, qna):
    messages = [
        {"role": "system", "content": f"""You are a professional Business Analyst writing content for a Business Requirements Document (BRD).
Use a formal and concise tone. Ensure clarity, no repetition, and structure the output with bullet points."""},
        {"role": "user", "content": f"""Using the following Q&A, write a well-organized and professional '{section}' section for the BRD.
Summarize insights, combine related points, and avoid repeating similar answers:
 
{qna}"""}
    ]
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=messages,
        temperature=0.2,
        max_tokens=4096
    )
    return response.choices[0].message.content
 
def save_brd_to_word(brd_sections_content):
    doc = Document()
    doc.add_heading("Business Requirement Document (BRD)", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\n")
    for section, content in brd_sections_content.items():
        doc.add_heading(section, level=2)
        for line in content.strip().split("\n"):
            if line.strip():
                doc.add_paragraph(f"• {line.strip()}")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
 
# === Streamlit UI ===
st.set_page_config(page_title="BRD Generator", layout="wide")
tabs = st.sidebar.radio("TABS", ["BRD Generation", "Design", "Testing", "Analysis"])
 
if tabs == "BRD Generation":
    st.title("📄 Business Requirement Document Generator")
 
    # === SOW or Manual Input ===
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        sow_file = st.file_uploader("Upload SOW", type=["pdf"])
    with col2:
        write_mode = st.checkbox("Write Requirements Manually")
    with col3:
        basic_qa = st.button("Basic Q/A")
 
    business_requirement = ""
    if sow_file:
        business_requirement = extract_text_from_pdf(sow_file)
    elif write_mode:
        business_requirement = st.text_area("Input Requirements")
 
    # === Summarize if too long ===
    if len(business_requirement) > 3000:
        st.warning("The business requirement is too long. A summarized version will be used.")
        summary_prompt = [
            {"role": "system", "content": "You are a summarizer for business requirement documents. Focus on key goals, scope, and technical requirements."},
            {"role": "user", "content": business_requirement}
        ]
        summary_response = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=summary_prompt,
            temperature=0.2,
            max_tokens=4096
        )
        business_requirement = summary_response.choices[0].message.content
 
    if basic_qa:
        st.subheader("Basic Questions")
        st.text_input("1. What is the objective of the project?")
        st.text_input("2. What is the scope of the work?")
        st.text_input("3. What are the technical constraints or requirements?")
        st.text_input("4. What data sources are involved?")
        st.text_input("5. What are the expected deliverables?")
 
    # === Generate Questions ===
    if st.button("Execute"):
        if not business_requirement:
            st.warning("Please upload a SOW or write requirements.")
        else:
            st.info("Generating questions... please wait ⌛")
            all_questions = {}
            for section, prompt in brd_sections.items():
                all_questions[section] = generate_section_questions(section, prompt, business_requirement)
            excel_data = save_questionnaire_to_excel(all_questions)
            st.download_button("Download Questionnaire Excel", data=excel_data, file_name="questionnaire.xlsx")
 
    # === Input Answers / MOM / Transcript ===
    st.subheader("Answer the Questions")
    qna_file = st.file_uploader("Upload Answered Questionnaire (Excel)", type=["xlsx"])
    transcript_files = st.file_uploader("Upload MOMs/Transcripts (DOCX)", type=["docx"], accept_multiple_files=True)
 
    # === Process Transcripts and Summarize if needed ===
    transcript_texts = []
    if transcript_files:
        for f in transcript_files:
            text = extract_text_from_docx(f)
            if len(text) > 3000:
                st.warning(f"Transcript '{f.name}' is too long. A summarized version will be used.")
                summary_prompt = [
                    {"role": "system", "content": "You are a BRD-focused summarizer. Extract all important discussion points from this meeting transcript."},
                    {"role": "user", "content": text}
                ]
                summary_response = client.chat.completions.create(
                    model="llama3-70b-8192",
                    messages=summary_prompt,
                    temperature=0.2,
                    max_tokens=4096
                )
                text = summary_response.choices[0].message.content
            transcript_texts.append(text)
 
    # === Generate BRD ===
    if st.button("Execute and Generate BRD"):
        if not qna_file:
            st.warning("Please upload the answered questionnaire Excel file.")
        else:
            df = pd.read_excel(qna_file)
            brd_final_sections = {}
            for section in brd_sections.keys():
                df_section = df[df['Section'] == section]
                qna = "\n".join(df_section.apply(lambda row: f"Q: {row['Questions']}\nA: {row['Answers']}", axis=1))
                brd_final_sections[section] = generate_brd_section(section, qna)
 
            # === Apply all transcript updates ===
            for transcript_text in transcript_texts:
                for section, content in brd_final_sections.items():
                    update_prompt = [
                        {"role": "system", "content": """You are a skilled Business Analyst updating a BRD section based on stakeholder meeting inputs.
Keep the tone professional and use clear bullet points."""},
                        {"role": "user", "content": f"""The following is an existing section titled '{section}' from a BRD:
{content}
 
Below is new transcript input from stakeholder meetings:
{transcript_text}
 
Please revise the section accordingly. Integrate any new insights, merge related points, and preserve the overall structure and tone. Use bullet points for clarity."""}
                    ]
                    response = client.chat.completions.create(
                        model="llama3-70b-8192",
                        messages=update_prompt,
                        temperature=0.2,
                        max_tokens=4096
                    )
                    brd_final_sections[section] = response.choices[0].message.content
 
            brd_file = save_brd_to_word(brd_final_sections)
            st.success("✅ BRD generated successfully!")
            st.download_button("🗕️ Download BRD (Word)", data=brd_file, file_name="Editable_BRD.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.button("Finalize BRD")
 
 